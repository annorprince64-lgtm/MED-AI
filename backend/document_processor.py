"""
Document Processor Module
Handles extraction and processing of various document formats:
- PDF (native text extraction)
- PDF (scanned documents via OCR)
- DOCX (Microsoft Word)
- Images (PNG, JPG, etc. via OCR)
- TXT (plain text)
- Tables extraction

Author: Annor Prince & Collins Yeboah
"""

import os
import io
import base64
import tempfile
import logging
from typing import Optional, Dict, List, Tuple, Any
from dataclasses import dataclass

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    text: str
    page_count: int
    file_type: str
    has_images: bool
    has_tables: bool
    metadata: Dict[str, Any]
    chunks: List[str]
    success: bool
    error: Optional[str] = None


class DocumentProcessor:
    """
    Main document processing class.
    Supports multiple formats with intelligent parsing.
    """
    
    # Supported file types
    SUPPORTED_TYPES = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'doc': ['application/msword'],
        'txt': ['text/plain'],
        'image': ['image/png', 'image/jpeg', 'image/jpg', 'image/gif', 'image/webp'],
    }
    
    # Maximum file sizes (in bytes)
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_TEXT_LENGTH = 100000  # Characters before chunking
    CHUNK_SIZE = 8000  # Characters per chunk
    CHUNK_OVERLAP = 500  # Overlap between chunks
    
    def __init__(self, enable_ocr: bool = True, enable_vision: bool = False):
        """
        Initialize the document processor.
        
        Args:
            enable_ocr: Enable OCR for scanned documents/images
            enable_vision: Enable vision AI for complex documents
        """
        self.enable_ocr = enable_ocr
        self.enable_vision = enable_vision
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and log available dependencies"""
        self.dependencies = {
            'pymupdf': False,
            'python_docx': False,
            'pillow': False,
            'pytesseract': False,
            'pdf2image': False,
            'camelot': False,
        }
        
        try:
            import fitz  # PyMuPDF
            self.dependencies['pymupdf'] = True
            logger.info("PyMuPDF available for PDF processing")
        except ImportError:
            logger.warning("PyMuPDF not installed - PDF support limited")
        
        try:
            from docx import Document
            self.dependencies['python_docx'] = True
            logger.info("python-docx available for DOCX processing")
        except ImportError:
            logger.warning("python-docx not installed - DOCX support unavailable")
        
        try:
            from PIL import Image
            self.dependencies['pillow'] = True
            logger.info("Pillow available for image processing")
        except ImportError:
            logger.warning("Pillow not installed - image support limited")
        
        try:
            import pytesseract
            self.dependencies['pytesseract'] = True
            logger.info("Tesseract available for OCR")
        except ImportError:
            logger.warning("pytesseract not installed - OCR unavailable")
        
        try:
            from pdf2image import convert_from_bytes
            self.dependencies['pdf2image'] = True
            logger.info("pdf2image available for PDF to image conversion")
        except ImportError:
            logger.warning("pdf2image not installed - scanned PDF support limited")
    
    def detect_file_type(self, file_data: bytes, filename: str = "") -> Tuple[str, str]:
        """
        Detect file type from content and filename.
        
        Returns:
            Tuple of (category, specific_type)
        """
        # Try to detect from magic bytes
        if file_data[:4] == b'%PDF':
            return ('pdf', 'application/pdf')
        
        if file_data[:4] == b'PK\x03\x04':
            # Could be DOCX (which is a ZIP file)
            if filename.lower().endswith('.docx'):
                return ('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            # Could be other ZIP-based formats
        
        if file_data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            # Old DOC format (OLE compound document)
            return ('doc', 'application/msword')
        
        # Check image magic bytes
        if file_data[:8] == b'\x89PNG\r\n\x1a\n':
            return ('image', 'image/png')
        if file_data[:2] == b'\xff\xd8':
            return ('image', 'image/jpeg')
        if file_data[:6] in [b'GIF87a', b'GIF89a']:
            return ('image', 'image/gif')
        if file_data[:4] == b'RIFF' and file_data[8:12] == b'WEBP':
            return ('image', 'image/webp')
        
        # Fallback to extension
        ext = filename.lower().split('.')[-1] if filename else ''
        if ext == 'pdf':
            return ('pdf', 'application/pdf')
        if ext == 'docx':
            return ('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if ext == 'doc':
            return ('doc', 'application/msword')
        if ext in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
            return ('image', f'image/{ext}')
        if ext == 'txt':
            return ('txt', 'text/plain')
        
        return ('unknown', 'application/octet-stream')
    
    def process(self, file_data: bytes, filename: str = "") -> ProcessedDocument:
        """
        Process a document and extract its text content.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename (used for type detection)
            
        Returns:
            ProcessedDocument with extracted text and metadata
        """
        # Check file size
        if len(file_data) > self.MAX_FILE_SIZE:
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="unknown",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"File too large. Maximum size is {self.MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Detect file type
        category, mime_type = self.detect_file_type(file_data, filename)
        
        logger.info(f"Processing {filename} as {category} ({mime_type})")
        
        # Route to appropriate processor
        try:
            if category == 'pdf':
                return self._process_pdf(file_data, filename)
            elif category == 'docx':
                return self._process_docx(file_data, filename)
            elif category == 'doc':
                return self._process_doc(file_data, filename)
            elif category == 'image':
                return self._process_image(file_data, filename)
            elif category == 'txt':
                return self._process_text(file_data, filename)
            else:
                # Try as plain text
                try:
                    text = file_data.decode('utf-8')
                    return self._create_result(text, 'text', 'txt')
                except:
                    return ProcessedDocument(
                        text="",
                        page_count=0,
                        file_type="unknown",
                        has_images=False,
                        has_tables=False,
                        metadata={},
                        chunks=[],
                        success=False,
                        error="Unsupported file format"
                    )
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type=category,
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=str(e)
            )
    
    def _process_pdf(self, file_data: bytes, filename: str) -> ProcessedDocument:
        """Process PDF files with multiple extraction methods"""
        
        if not self.dependencies['pymupdf']:
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="pdf",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error="PDF processing unavailable. Install PyMuPDF: pip install pymupdf"
            )
        
        import fitz  # PyMuPDF
        
        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
            page_count = len(doc)
            
            # Extract metadata
            metadata = {
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'page_count': page_count,
            }
            
            # Try text extraction first
            all_text = []
            has_images = False
            has_tables = False
            pages_with_no_text = []
            
            for page_num, page in enumerate(doc):
                # Extract text
                text = page.get_text("text")
                
                # Check for images
                images = page.get_images()
                if images:
                    has_images = True
                
                # Check for tables (heuristic: look for grid-like structures)
                tables = page.find_tables()
                if tables and len(tables.tables) > 0:
                    has_tables = True
                    # Extract table text separately
                    for table in tables.tables:
                        table_text = table.extract()
                        if table_text:
                            # Format table as markdown
                            formatted_table = self._format_table(table_text)
                            text += f"\n\n[TABLE]\n{formatted_table}\n[/TABLE]\n"
                
                if text.strip():
                    all_text.append(f"--- Page {page_num + 1} ---\n{text}")
                else:
                    pages_with_no_text.append(page_num)
            
            combined_text = "\n\n".join(all_text)
            
            # If no text extracted, try OCR
            if not combined_text.strip() and pages_with_no_text and self.enable_ocr:
                logger.info(f"No text found, attempting OCR on {len(pages_with_no_text)} pages")
                ocr_text = self._ocr_pdf_pages(file_data, pages_with_no_text)
                if ocr_text:
                    combined_text = ocr_text
            
            doc.close()
            
            return self._create_result(
                combined_text,
                'pdf',
                'pdf',
                page_count=page_count,
                has_images=has_images,
                has_tables=has_tables,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="pdf",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"Failed to process PDF: {str(e)}"
            )
    
    def _ocr_pdf_pages(self, file_data: bytes, page_numbers: List[int] = None) -> str:
        """Perform OCR on PDF pages"""
        if not self.dependencies['pdf2image'] or not self.dependencies['pytesseract']:
            return ""
        
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            
            # Convert PDF to images
            images = convert_from_bytes(file_data, dpi=200)
            
            if page_numbers:
                images = [img for i, img in enumerate(images) if i in page_numbers]
            
            ocr_results = []
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image, lang='eng')
                if text.strip():
                    ocr_results.append(f"--- Page {i + 1} (OCR) ---\n{text}")
            
            return "\n\n".join(ocr_results)
            
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            return ""
    
    def _process_docx(self, file_data: bytes, filename: str) -> ProcessedDocument:
        """Process DOCX files"""
        if not self.dependencies['python_docx']:
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="docx",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error="DOCX processing unavailable. Install python-docx: pip install python-docx"
            )
        
        from docx import Document
        
        try:
            doc = Document(io.BytesIO(file_data))
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check for heading styles
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        paragraphs.append(f"\n{'#' * int(level) if level.isdigit() else '##'} {para.text}\n")
                    else:
                        paragraphs.append(para.text)
            
            # Extract tables
            has_tables = len(doc.tables) > 0
            if has_tables:
                paragraphs.append("\n\n--- Tables ---\n")
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    formatted = self._format_table(table_data)
                    paragraphs.append(f"\n{formatted}\n")
            
            combined_text = "\n".join(paragraphs)
            
            return self._create_result(
                combined_text,
                'docx',
                'docx',
                has_images=False,  # We don't extract embedded images yet
                has_tables=has_tables,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="docx",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"Failed to process DOCX: {str(e)}"
            )
    
    def _process_doc(self, file_data: bytes, filename: str) -> ProcessedDocument:
        """Process old DOC files (limited support)"""
        # Old .doc format is harder to parse
        # Best approach is to use antiword or convert to docx
        return ProcessedDocument(
            text="",
            page_count=0,
            file_type="doc",
            has_images=False,
            has_tables=False,
            metadata={},
            chunks=[],
            success=False,
            error="Old .doc format not supported. Please convert to .docx format."
        )
    
    def _process_image(self, file_data: bytes, filename: str) -> ProcessedDocument:
        """Process image files using OCR"""
        if not self.dependencies['pillow']:
            return ProcessedDocument(
                text="",
                page_count=1,
                file_type="image",
                has_images=True,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error="Image processing unavailable. Install Pillow: pip install Pillow"
            )
        
        from PIL import Image
        
        try:
            image = Image.open(io.BytesIO(file_data))
            
            metadata = {
                'format': image.format,
                'size': image.size,
                'mode': image.mode,
            }
            
            # Perform OCR
            extracted_text = ""
            if self.enable_ocr and self.dependencies['pytesseract']:
                import pytesseract
                extracted_text = pytesseract.image_to_string(image, lang='eng')
            
            if not extracted_text.strip():
                extracted_text = "[Image file - no text detected]"
            
            return self._create_result(
                extracted_text,
                'image',
                'image',
                page_count=1,
                has_images=True,
                has_tables=False,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Image processing error: {str(e)}")
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="image",
                has_images=True,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"Failed to process image: {str(e)}"
            )
    
    def _process_text(self, file_data: bytes, filename: str) -> ProcessedDocument:
        """Process plain text files"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = file_data.decode('utf-8', errors='replace')
            
            return self._create_result(text, 'text', 'txt')
            
        except Exception as e:
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="text",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"Failed to process text file: {str(e)}"
            )
    
    def _format_table(self, table_data: List[List[str]]) -> str:
        """Format table data as markdown"""
        if not table_data:
            return ""
        
        # Determine column widths
        num_cols = max(len(row) for row in table_data)
        col_widths = [0] * num_cols
        
        for row in table_data:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(str(cell)))
        
        # Format as markdown table
        lines = []
        
        # Header row
        if table_data:
            header = table_data[0]
            header_line = "| " + " | ".join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(header)) + " |"
            lines.append(header_line)
            
            # Separator
            sep_line = "|" + "|".join("-" * (w + 2) for w in col_widths) + "|"
            lines.append(sep_line)
            
            # Data rows
            for row in table_data[1:]:
                row_line = "| " + " | ".join(str(cell).ljust(col_widths[i]) if i < len(row) else "" for i in range(num_cols)) + " |"
                lines.append(row_line)
        
        return "\n".join(lines)
    
    def _create_result(
        self,
        text: str,
        file_type: str,
        extension: str,
        page_count: int = 1,
        has_images: bool = False,
        has_tables: bool = False,
        metadata: Dict = None
    ) -> ProcessedDocument:
        """Create the final ProcessedDocument with chunking"""
        
        metadata = metadata or {}
        
        # Create chunks for large texts
        chunks = self._create_chunks(text) if len(text) > self.MAX_TEXT_LENGTH else [text]
        
        return ProcessedDocument(
            text=text,
            page_count=page_count,
            file_type=file_type,
            has_images=has_images,
            has_tables=has_tables,
            metadata=metadata,
            chunks=chunks,
            success=True
        )
    
    def _create_chunks(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for processing.
        Tries to split at sentence/paragraph boundaries.
        """
        if len(text) <= self.CHUNK_SIZE:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.CHUNK_SIZE
            
            # Try to find a good break point
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + self.CHUNK_SIZE // 2:
                    end = para_break + 2
                else:
                    # Look for sentence break
                    sentence_break = text.rfind('. ', start, end)
                    if sentence_break > start + self.CHUNK_SIZE // 2:
                        end = sentence_break + 2
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - self.CHUNK_OVERLAP if end < len(text) else len(text)
        
        return chunks
    
    def process_base64(self, base64_data: str, filename: str = "") -> ProcessedDocument:
        """
        Process a document from base64 encoded data.
        
        Args:
            base64_data: Base64 encoded file data (with or without data URI prefix)
            filename: Original filename
            
        Returns:
            ProcessedDocument
        """
        try:
            # Remove data URI prefix if present
            if ',' in base64_data:
                base64_data = base64_data.split(',', 1)[1]
            
            file_data = base64.b64decode(base64_data)
            return self.process(file_data, filename)
            
        except Exception as e:
            logger.error(f"Base64 decoding error: {str(e)}")
            return ProcessedDocument(
                text="",
                page_count=0,
                file_type="unknown",
                has_images=False,
                has_tables=False,
                metadata={},
                chunks=[],
                success=False,
                error=f"Failed to decode file data: {str(e)}"
            )


# Convenience function for direct use
def process_document(file_data: bytes, filename: str = "", enable_ocr: bool = True) -> ProcessedDocument:
    """
    Convenience function to process a document.
    
    Args:
        file_data: Raw file bytes
        filename: Original filename
        enable_ocr: Whether to enable OCR for scanned content
        
    Returns:
        ProcessedDocument with extracted text
    """
    processor = DocumentProcessor(enable_ocr=enable_ocr)
    return processor.process(file_data, filename)


def process_document_base64(base64_data: str, filename: str = "", enable_ocr: bool = True) -> ProcessedDocument:
    """
    Convenience function to process a document from base64 data.
    
    Args:
        base64_data: Base64 encoded file data
        filename: Original filename
        enable_ocr: Whether to enable OCR for scanned content
        
    Returns:
        ProcessedDocument with extracted text
    """
    processor = DocumentProcessor(enable_ocr=enable_ocr)
    return processor.process_base64(base64_data, filename)
